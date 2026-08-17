# -*- coding: utf-8 -*-
"""
generate_references_docx.py - 生成"引用与借鉴记录"Word 文档
用途：单独记录法律文书所引用的法条、案例、数据来源与核验状态，方便查看与追溯。
用法： python generate_references_docx.py
输出：E:\\Aether\\projects\\legal-skill\\output\\引用与借鉴记录-<日期>.docx
"""
import json, os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

ROOT = os.path.join(os.path.dirname(__file__), "..")
LAWS = os.path.join(ROOT, "knowledge", "laws", "laws.json")
LPR = os.path.join(ROOT, "knowledge", "lpr", "lpr.json")
OUT_DIR = os.path.join(ROOT, "output")


def set_font(run, name="仿宋", size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), "宋体", 16, True)
    p.paragraph_format.space_after = Pt(14)


def add_h(doc, text):
    p = doc.add_paragraph()
    set_font(p.add_run(text), "黑体", 13, True)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)


def add_p(doc, text, indent=True, size=12, bold=False):
    p = doc.add_paragraph()
    set_font(p.add_run(text), "仿宋", size, bold)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.5) if indent else Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def main():
    with open(LAWS, "r", encoding="utf-8") as f:
        laws_kb = json.load(f)
    with open(LPR, "r", encoding="utf-8") as f:
        lpr_kb = json.load(f)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(3.0)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.6)

    add_title(doc, "引用与借鉴记录")

    add_p(doc, f"生成日期：{date.today().isoformat()}｜ 用途：记录法律文书所引用的法条、案例、数据来源与核验状态，供查看与追溯。", indent=False, size=11)

    # 一、法条引用
    add_h(doc, "一、法条引用记录")
    for lk, law in laws_kb["laws"].items():
        add_p(doc, f"【{law['name']}】（{law['status']}）", indent=False, bold=True)
        for art in law["articles"]:
            add_p(doc, f"□ {art['article']}　核验状态：{art['verify']}　｜　{art['content'][:60]}{'…' if len(art['content'])>60 else ''}")
    add_p(doc, "核验入口：国家法律法规数据库 https://flk.npc.gov.cn", indent=False, size=11)

    # 二、案例引用
    add_h(doc, "二、案例引用记录（真实可核验）")
    cases = [
        ("(2014)民一终字第38号", "李占江、朱丽敏与贝洪峰、沈阳东昊地产有限公司民间借贷纠纷案（《最高人民法院公报》2015年第9期）",
         "本金以实际交付为准；利息超「银行同类贷款利率四倍」部分不予保护。适用1991年「四倍」规则，引用须注明时点。",
         "http://gongbao.court.gov.cn/Details/8e5a855a1d656c6b26158702e72098.html", "已核验"),
        ("(2019)最高法民终218号", "黄明与陈琪玲、陈泽峰、福建省丰泉环保集团有限公司民间借贷纠纷案（《最高人民法院公报》2022年第6期）",
         "约定利率超36%部分无效，已付利息按36%折抵本金，此后按24%计算。适用2015年24%/36%规则，引用须注明时点。",
         "http://gongbao.court.gov.cn/Details/06f1d4f912cdfd7d53d0d42c0dbc5c.html", "已核验"),
        ("(2015)民一终字第260号", "邵萍与云南通海昆通工贸有限公司、通海兴通达工贸有限公司民间借贷纠纷案（《最高人民法院公报》2017年第3期）",
         "公司人格混同连带责任，借款主体认定参考。依据《公司法》第二十条第三款。",
         "http://gongbao.court.gov.cn/Details/dd50d3d7d1522906721893219a1465.html", "已核验"),
    ]
    for cno, cname, cpoint, curl, cverify in cases:
        add_p(doc, f"案号：{cno}", indent=False, bold=True)
        add_p(doc, f"案件：{cname}")
        add_p(doc, f"裁判要旨：{cpoint}")
        add_p(doc, f"来源：{curl}")
        add_p(doc, f"核验状态：{cverify}")

    # 三、真实但无案号判例
    add_h(doc, "三、真实判例（无公开案号，仅可作参考，不作案号引用）")
    refs = [
        "房山法院（北京法院网 2023-05-17）：出借信用卡+贷款转贷→合同无效，仅返还本金。 https://bjgy.bjcourt.gov.cn/article/detail/2023/05/id/7295776.shtml",
        "延庆法院（2025-02-18）：银行借款转贷→合同无效，按LPR计息，不支持约定15%。 https://bjgy.bjcourt.gov.cn/article/detail/2025/02/id/8710897.shtml",
        "沛县法院（江苏法院网 2026-04-13）：两年向7人出借超400万→职业放贷，合同无效，仅返本金。 http://www.jsfy.gov.cn/article/107311.html",
        "邳州法院（2024-08-01）：两年起诉16件→合同无效，返本金按LPR付占用费。 http://www.jsfy.gov.cn/article/99848.html",
        "北京高院典型案例1（2023-12-26）：仅凭转账凭证、不能证明借贷合意→不予支持。 https://bjgy.bjcourt.gov.cn/article/detail/2023/12/id/7727777.shtml",
        "珲春市法院砍头息案（人民法院报 2026-05-12）：借条载20万实付18万→按实际交付认定本金。 https://www.rmfyb.com/content/202605/12/article_1024014_1391811899_6606786.html",
    ]
    for r in refs:
        add_p(doc, f"● {r}")
    add_p(doc, "⚠️ 上述判例均真实存在（官方渠道），但公开报道未载案号，不可当作案号引用；如需案号须在人民法院案例库登录检索。", indent=False, bold=True)

    # 四、尚未取得可核验案号的点
    add_h(doc, "四、尚未取得可核验案号的点（诚实记录）")
    for pt in ["现行 LPR 四倍规则（2020修正第25条）直接适用案例", "逾期利息+违约金合计≤LPR四倍（第29条）"]:
        add_p(doc, f"□ {pt}：无可核验案号（人民法院案例库已强制登录），引用时用「该类型已有生效判例（案号可在人民法院案例库检索）」。")
    # 五、数据引用
    add_h(doc, "五、数据引用记录")
    cur = lpr_kb.get("current", {})
    add_p(doc, f"LPR（截至 {cur.get('as_of','待更新')}）：1年期 {cur.get('lpr_1y')}，5年期以上 {cur.get('lpr_5y')}；民间借贷利率司法保护上限 = {cur.get('private_lending_cap')}。")
    add_p(doc, f"来源：中国人民银行授权全国银行间同业拆借中心公告 {cur.get('source_url','chinamoney.com.cn')}")
    add_p(doc, "更新方式：每月20日运行 scripts\\fetch_lpr.py 自动抓取。")

    # 六、模板依据
    add_h(doc, "六、文书模板依据")
    add_p(doc, "《最高人民法院、司法部、中华全国律师协会关于印发部分案件民事起诉状、答辩状示范文本（试行）的通知》（2024-03-04，11类44份，含民间借贷）。")

    # 七、核验状态汇总
    add_h(doc, "七、核验状态汇总")
    verified = sum(1 for law in laws_kb["laws"].values() for a in law["articles"] if a["verify"] == "已核验")
    pending = sum(1 for law in laws_kb["laws"].values() for a in law["articles"] if a["verify"] != "已核验")
    add_p(doc, f"法条：已核验 {verified} 条 / 待复核 {pending} 条（待复核条目见 knowledge\\laws\\laws.json，用 scripts\\fetch_law.py 管理）。")
    add_p(doc, f"案例：可核验案号 3 个（全部已核验），无案号真实判例 6 条，待补案号 2 点。")

    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, f"引用与借鉴记录-{date.today().isoformat()}.docx")
    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
