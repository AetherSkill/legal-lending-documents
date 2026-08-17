# -*- coding: utf-8 -*-
"""
generate_docx.py - 生成排版规范的民间借贷起诉状 Word 文档
用法：
  python generate_docx.py                            # 用内置测试案例
  python generate_docx.py --case case.json           # 用外部案例 JSON
输出：E:\\Aether\\projects\\legal-skill\\output\\民事起诉状-<日期>.docx
"""
import argparse, json, os, sys
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "output")

DEFAULT_CASE = {
    "plaintiff": {"name": "张三", "gender": "男", "birth": "____", "ethnic": "汉",
                  "occupation": "____", "address": "__________", "id": "__________", "phone": "__________"},
    "defendant": {"name": "李四", "gender": "男", "birth": "____", "ethnic": "汉",
                  "occupation": "____", "address": "__________", "phone": "__________"},
    "principal": 50000,
    "lend_date": "2023年3月1日",
    "delivery": "银行转账",
    "agreed_rate": "15%",
    "claim_rate": "14.6%",
    "due_date": "2023年9月1日",
    "repaid_principal": 0,
    "repaid_interest": 0,
    "interest_tmp": "25,280",
    "lpr_cap": "14.6% (合同成立时1年期LPR 3.65% × 4)",
    "evidence": [
        "银行转账凭证（证明借款实际交付——自然人借贷自交付时成立，民法典第六百七十九条）",
        "借条/借款合同（证明借贷合意及约定内容）",
        "微信聊天记录/短信/通话录音（证明催收事实）",
        "原、被告身份证复印件（主体资格）",
    ],
}


def set_run_font(run, name_cn="仿宋", size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), name_cn)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, name_cn="宋体", size=16, bold=True)
    p.paragraph_format.space_after = Pt(18)


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, name_cn="黑体", size=13, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, name_cn="仿宋", size=12)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.06) if indent else Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    return p


def build(case):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(3.7); sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.6)

    add_title(doc, "民事起诉状（民间借贷纠纷）")

    add_heading(doc, "原告信息")
    pi, di = case["plaintiff"], case["defendant"]
    add_body(doc, f"原告：{pi['name']}，{pi['gender']}，{pi['birth']}出生，{pi['ethnic']}族，{pi['occupation']}，住{pi['address']}，身份证号：{pi['id']}，联系电话：{pi['phone']}。")
    add_body(doc, f"被告：{di['name']}，{di['gender']}，{di['birth']}出生，{di['ethnic']}族，{di['occupation']}，住{di['address']}，联系电话：{di['phone']}。")

    add_heading(doc, "一、诉讼请求")
    add_body(doc, f"1. 判令被告向原告偿还借款本金人民币 {case['principal']:,} 元；")
    add_body(doc, f"2. 判令被告支付利息（以 {case['principal']:,} 元为基数，自 {case['lend_date']} 起按年利率 {case['claim_rate']} 计算至实际清偿之日止，暂计至起诉之日为 {case['interest_tmp']} 元）；")
    add_body(doc, "3. 判令被告支付逾期利息（如与其他费用合计超过合同成立时一年期贷款市场报价利率四倍的，以四倍为限）；")
    add_body(doc, "4. 本案诉讼费、保全费由被告负担。")

    add_heading(doc, "二、事实与理由（要素式）")
    add_body(doc, f"□ 借款本金金额：{case['principal']:,} 元；币种：人民币")
    add_body(doc, f"□ 借款日期：{case['lend_date']}")
    add_body(doc, f"□ 交付方式：{case['delivery']}（附交付凭证）")
    add_body(doc, f"□ 约定利率：年利率 {case['agreed_rate']}（已提示：超过司法保护上限部分不予支持，按 {case['claim_rate']} 主张）")
    add_body(doc, f"□ 约定还款日期：{case['due_date']}")
    add_body(doc, "□ 逾期利率/违约金：无约定（主张按法定上限计算逾期利息）")
    add_body(doc, f"□ 已还金额：本金 {case['repaid_principal']} 元，利息 {case['repaid_interest']} 元")
    add_body(doc, "□ 合同效力排查：无套取转贷、无职业放贷、无违法犯罪用途、无违背公序良俗情形")
    add_body(doc, "理由摘要：被告经多次催收仍未履行还款义务，已构成违约。依据《中华人民共和国民法典》第五百七十九条、第六百七十六条，《最高人民法院关于审理民间借贷案件适用法律若干问题的规定》第二十五条第一款、第二十九条，被告应偿还本金并支付利息。")

    add_heading(doc, "三、证据清单")
    for i, ev in enumerate(case["evidence"], 1):
        add_body(doc, f"{i}. {ev}")

    add_heading(doc, "四、此致")
    add_body(doc, "此致")
    add_body(doc, "____________人民法院")
    doc.add_paragraph()
    add_body(doc, "附：起诉状副本 1 份；证据副本 1 份。", indent=False)
    add_body(doc, f"原告（签名/盖章）：{case['plaintiff']['name']}", indent=False)
    add_body(doc, f"{date.today().year}年__月__日", indent=False)

    add_heading(doc, "附注")
    add_body(doc, f"司法保护利率上限：{case['lpr_cap']}。", indent=False)
    add_body(doc, "时效水印：本文书引用的 LPR 以借款成立时数值计算；现行 LPR 须于生成日经 chinamoney.com.cn 复核。法条以生成日 flk.npc.gov.cn 现行有效文本为准。", indent=False)
    add_body(doc, "免责声明：本文书由 AI 生成，仅供参考与初稿使用，不构成法律意见，不承诺诉讼结果；立案与正式提交前请以受诉法院要求为准并建议咨询执业律师。", indent=False)

    return doc


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--case", default=None)
    args = ap.parse_args()
    case = DEFAULT_CASE
    if args.case:
        with open(args.case, "r", encoding="utf-8") as f:
            case = json.load(f)
    doc = build(case)
    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, f"民事起诉状-民间借贷-{date.today().isoformat()}.docx")
    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
